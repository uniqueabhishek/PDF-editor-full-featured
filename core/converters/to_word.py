"""
Ultra PDF Editor - PDF to Word Converter
Convert PDF documents to Microsoft Word format
"""
from pathlib import Path
from typing import Union, Optional, List, Dict, Any
import fitz
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import io
import re


class PDFToWordConverter:
    """Converts PDF documents to Word (.docx) format"""

    def __init__(self, pdf_path: Union[str, Path] = None, document: fitz.Document = None):
        """
        Initialize converter

        Args:
            pdf_path: Path to PDF file
            document: Open fitz.Document
        """
        self._pdf_doc: Optional[fitz.Document] = None
        self._owns_doc = False

        if document:
            self._pdf_doc = document
        elif pdf_path:
            self._pdf_doc = fitz.open(str(pdf_path))
            self._owns_doc = True

    def __del__(self):
        if self._owns_doc and self._pdf_doc:
            self._pdf_doc.close()

    def set_document(self, document: fitz.Document):
        """Set the document to convert"""
        if self._owns_doc and self._pdf_doc:
            self._pdf_doc.close()
        self._pdf_doc = document
        self._owns_doc = False

    def convert(
        self,
        output_path: Union[str, Path],
        pages: Optional[List[int]] = None,
        include_images: bool = True,
        preserve_layout: bool = True,
        callback=None
    ) -> Path:
        """
        Convert PDF to Word document

        Args:
            output_path: Output .docx file path
            pages: List of page numbers to convert (None = all)
            include_images: Include images from PDF
            preserve_layout: Try to preserve original layout
            callback: Progress callback(current_page, total_pages)

        Returns:
            Path to output file
        """
        if not self._pdf_doc:
            raise ValueError("No document loaded")

        output_path = Path(output_path)

        if pages is None:
            pages = list(range(len(self._pdf_doc)))

        # Create Word document
        doc = Document()

        total_pages = len(pages)

        for i, page_num in enumerate(pages):
            page = self._pdf_doc[page_num]

            # Extract and add content
            if preserve_layout:
                self._convert_page_with_layout(doc, page, include_images)
            else:
                self._convert_page_simple(doc, page, include_images)

            # Add page break between pages (except last)
            if i < total_pages - 1:
                doc.add_page_break()

            if callback:
                callback(i + 1, total_pages)

        # Save document
        doc.save(str(output_path))
        return output_path

    def _convert_page_simple(self, doc: Document, page: fitz.Page, include_images: bool):
        """Simple text extraction without layout preservation"""
        # Extract text
        text = page.get_text("text")

        # Add text blocks
        for paragraph_text in text.split('\n\n'):
            paragraph_text = paragraph_text.strip()
            if paragraph_text:
                para = doc.add_paragraph(paragraph_text)

        # Extract and add images
        if include_images:
            self._add_page_images(doc, page)

    def _convert_page_with_layout(self, doc: Document, page: fitz.Page, include_images: bool):
        """Convert page while trying to preserve layout"""
        # Get text blocks with positioning
        blocks = page.get_text("dict")["blocks"]

        for block in blocks:
            if block["type"] == 0:  # Text block
                self._process_text_block(doc, block)
            elif block["type"] == 1 and include_images:  # Image block
                self._process_image_block(doc, block, page)

    def _process_text_block(self, doc: Document, block: Dict[str, Any]):
        """Process a text block from PDF"""
        for line in block.get("lines", []):
            line_text = ""
            font_size = 12
            is_bold = False
            is_italic = False

            for span in line.get("spans", []):
                text = span.get("text", "")
                line_text += text

                # Get font info
                font_size = span.get("size", 12)
                flags = span.get("flags", 0)
                is_bold = flags & 2 ** 4  # Bold flag
                is_italic = flags & 2 ** 1  # Italic flag

            if line_text.strip():
                para = doc.add_paragraph()
                run = para.add_run(line_text)
                run.font.size = Pt(font_size)
                run.bold = bool(is_bold)
                run.italic = bool(is_italic)

    def _process_image_block(self, doc: Document, block: Dict[str, Any], page: fitz.Page):
        """Process an image block from PDF"""
        # Get image from block
        bbox = block.get("bbox", (0, 0, 100, 100))
        width = bbox[2] - bbox[0]
        height = bbox[3] - bbox[1]

        # Try to extract the image
        try:
            # Get images on page
            images = page.get_images()
            if images:
                xref = images[0][0]
                image_info = self._pdf_doc.extract_image(xref)
                image_bytes = image_info["image"]

                # Add to document
                image_stream = io.BytesIO(image_bytes)
                # Scale image to reasonable size
                max_width = Inches(6)
                doc.add_picture(image_stream, width=min(Inches(width / 72), max_width))
        except Exception:
            pass  # Skip images that can't be extracted

    def _add_page_images(self, doc: Document, page: fitz.Page):
        """Add all images from a page"""
        images = page.get_images()

        for img_info in images:
            try:
                xref = img_info[0]
                image_data = self._pdf_doc.extract_image(xref)
                image_bytes = image_data["image"]

                image_stream = io.BytesIO(image_bytes)
                doc.add_picture(image_stream, width=Inches(5))
            except Exception:
                continue

    def extract_text(self, pages: Optional[List[int]] = None) -> str:
        """
        Extract plain text from PDF

        Args:
            pages: List of page numbers (None = all)

        Returns:
            Extracted text
        """
        if not self._pdf_doc:
            raise ValueError("No document loaded")

        if pages is None:
            pages = list(range(len(self._pdf_doc)))

        text_parts = []
        for page_num in pages:
            page = self._pdf_doc[page_num]
            text_parts.append(page.get_text("text"))

        return "\n\n".join(text_parts)

    def extract_tables(self, page_num: int) -> List[List[List[str]]]:
        """
        Extract tables from a page (basic implementation)

        Returns:
            List of tables, each table is a list of rows,
            each row is a list of cell values
        """
        if not self._pdf_doc:
            raise ValueError("No document loaded")

        # This is a simplified implementation
        # For better table extraction, consider using camelot or pdfplumber

        page = self._pdf_doc[page_num]
        tables = []

        # Try to find table-like structures in text blocks
        blocks = page.get_text("dict")["blocks"]

        # Group blocks by vertical position (potential rows)
        # This is a very basic approach

        return tables

    @property
    def page_count(self) -> int:
        """Get number of pages"""
        return len(self._pdf_doc) if self._pdf_doc else 0


def convert_pdf_to_word(
    pdf_path: Union[str, Path],
    output_path: Union[str, Path] = None,
    include_images: bool = True
) -> Path:
    """
    Convenience function to convert PDF to Word

    Args:
        pdf_path: Input PDF path
        output_path: Output path (defaults to same name with .docx)
        include_images: Include images in output

    Returns:
        Path to output file
    """
    pdf_path = Path(pdf_path)

    if output_path is None:
        output_path = pdf_path.with_suffix('.docx')

    converter = PDFToWordConverter(pdf_path)
    return converter.convert(output_path, include_images=include_images)


def convert_pdf_to_text(pdf_path: Union[str, Path], output_path: Union[str, Path] = None) -> Path:
    """
    Convert PDF to plain text file

    Args:
        pdf_path: Input PDF path
        output_path: Output path (defaults to same name with .txt)

    Returns:
        Path to output file
    """
    pdf_path = Path(pdf_path)

    if output_path is None:
        output_path = pdf_path.with_suffix('.txt')

    converter = PDFToWordConverter(pdf_path)
    text = converter.extract_text()

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(text)

    return output_path
