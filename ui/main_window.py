"""
Ultra PDF Editor - Main Application Window
"""
from PyQt6.QtWidgets import (
    QMainWindow, QWidget, QHBoxLayout, QSplitter,
    QStatusBar, QFileDialog, QMessageBox, QInputDialog,
    QProgressDialog, QApplication, QLabel
)
from PyQt6.QtCore import Qt
from PyQt6.QtGui import QAction, QKeySequence, QCloseEvent, QImage, QPixmap, QPageLayout
import fitz
from pathlib import Path
from typing import Optional, List, Dict, Any
import os
import datetime

from .pdf_viewer import PDFViewer, ToolMode, ViewMode
from .sidebar import Sidebar
from .toolbar import MainToolbar, AnnotationToolbar
from .dialogs import (
    FindDialog, FindReplaceDialog, ExtractPagesDialog, CropDialog,
    HeaderFooterDialog, BatchDialog
)
from core.pdf_document import PDFDocument
from config import config, UserSettings
from utils.history import HistoryManager, PageAddCommand, PageDeleteCommand, PageRotateCommand, AnnotationAddCommand


class MainWindow(QMainWindow):
    """Main application window"""

    def _create_action(self, text: str, slot=None, shortcut=None, checkable=False) -> QAction:
        """Helper to create menu actions (PyQt6 compatible)"""
        action = QAction(text, self)
        if slot:
            action.triggered.connect(slot)
        if shortcut:
            action.setShortcut(shortcut)
        if checkable:
            action.setCheckable(True)
        return action

    def __init__(self):
        super().__init__()

        # Core state
        self._document = PDFDocument()
        self._settings = UserSettings.load(config.SETTINGS_PATH)
        self._is_modified = False
        self._current_file: Optional[Path] = None

        # Undo/Redo history manager
        self._history_manager = HistoryManager(config.UNDO_HISTORY_SIZE)

        # Search state
        self._search_results: List[Dict] = []
        self._current_search_index = 0
        self._find_dialog: Optional[FindDialog] = None
        self._replace_dialog: Optional[FindReplaceDialog] = None

        # Setup UI
        self._setup_ui()
        self._setup_menus()
        self._setup_toolbars()
        self._setup_statusbar()
        self._connect_signals()
        self._apply_settings()

        # Set initial state
        self._update_title()
        self._update_actions_state()

    def _setup_ui(self):
        """Setup the main UI layout"""
        self.setWindowTitle(config.APP_NAME)
        self.setMinimumSize(config.WINDOW_MIN_WIDTH, config.WINDOW_MIN_HEIGHT)
        self.resize(config.WINDOW_WIDTH, config.WINDOW_HEIGHT)

        # Central widget with splitter
        central = QWidget()
        self.setCentralWidget(central)

        layout = QHBoxLayout(central)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)

        # Splitter for sidebar and viewer
        self._splitter = QSplitter(Qt.Orientation.Horizontal)

        # Sidebar
        self._sidebar = Sidebar()
        self._sidebar.setMinimumWidth(150)
        self._sidebar.setMaximumWidth(400)
        self._splitter.addWidget(self._sidebar)

        # PDF Viewer
        self._viewer = PDFViewer()
        self._splitter.addWidget(self._viewer)

        # Set initial sizes
        self._splitter.setSizes(
            [config.SIDEBAR_WIDTH, config.WINDOW_WIDTH - config.SIDEBAR_WIDTH])

        layout.addWidget(self._splitter)

    def _setup_menus(self):
        """Setup menu bar"""
        menubar = self.menuBar()
        assert menubar is not None  # QMainWindow always has a menu bar

        # === File Menu ===
        file_menu = menubar.addMenu("&File")
        assert file_menu is not None

        self._new_action = file_menu.addAction("&New", self._new_document)
        self._new_action.setShortcut(QKeySequence.StandardKey.New)

        self._open_action = file_menu.addAction(
            "&Open...", self._open_document)
        self._open_action.setShortcut(QKeySequence.StandardKey.Open)

        # Recent files submenu
        self._recent_menu = file_menu.addMenu("Recent Files")
        assert self._recent_menu is not None
        self._update_recent_files_menu()

        file_menu.addSeparator()

        self._save_action = file_menu.addAction("&Save", self._save_document)
        self._save_action.setShortcut(QKeySequence.StandardKey.Save)

        self._save_as_action = file_menu.addAction(
            "Save &As...", self._save_document_as)
        self._save_as_action.setShortcut(QKeySequence("Ctrl+Shift+S"))

        file_menu.addSeparator()

        self._close_action = file_menu.addAction(
            "&Close", self._close_document)
        self._close_action.setShortcut(QKeySequence.StandardKey.Close)

        file_menu.addSeparator()

        # Export submenu
        export_menu = file_menu.addMenu("Export")
        assert export_menu is not None
        export_menu.addAction("Export as Images...", self._export_as_images)
        export_menu.addAction("Export as Word...", self._export_as_word)
        export_menu.addAction("Export as Text...", self._export_as_text)

        file_menu.addSeparator()

        self._print_action = file_menu.addAction(
            "&Print...", self._print_document)
        self._print_action.setShortcut(QKeySequence.StandardKey.Print)

        file_menu.addSeparator()

        self._properties_action = file_menu.addAction(
            "Properties...", self._show_properties)

        file_menu.addSeparator()

        file_menu.addAction("E&xit", self.close)

        # === Edit Menu ===
        edit_menu = menubar.addMenu("&Edit")
        assert edit_menu is not None

        self._undo_action = edit_menu.addAction("&Undo", self._undo)
        self._undo_action.setShortcut(QKeySequence.StandardKey.Undo)

        self._redo_action = edit_menu.addAction("&Redo", self._redo)
        self._redo_action.setShortcut(QKeySequence.StandardKey.Redo)

        edit_menu.addSeparator()

        self._cut_action = edit_menu.addAction("Cu&t", self._cut)
        self._cut_action.setShortcut(QKeySequence.StandardKey.Cut)

        self._copy_action = edit_menu.addAction("&Copy", self._copy)
        self._copy_action.setShortcut(QKeySequence.StandardKey.Copy)

        self._paste_action = edit_menu.addAction("&Paste", self._paste)
        self._paste_action.setShortcut(QKeySequence.StandardKey.Paste)

        self._delete_action = edit_menu.addAction("&Delete", self._delete)
        self._delete_action.setShortcut(QKeySequence.StandardKey.Delete)

        edit_menu.addSeparator()

        self._select_all_action = edit_menu.addAction(
            "Select &All", self._select_all)
        self._select_all_action.setShortcut(QKeySequence.StandardKey.SelectAll)

        edit_menu.addSeparator()

        self._find_action = edit_menu.addAction("&Find...", self._show_find)
        self._find_action.setShortcut(QKeySequence.StandardKey.Find)

        self._replace_action = edit_menu.addAction(
            "Find && &Replace...", self._show_replace)
        self._replace_action.setShortcut(QKeySequence("Ctrl+H"))

        # === View Menu ===
        view_menu = menubar.addMenu("&View")
        assert view_menu is not None

        # Zoom submenu
        zoom_menu = view_menu.addMenu("Zoom")
        assert zoom_menu is not None
        zoom_in_action = zoom_menu.addAction("Zoom In")
        zoom_in_action.setShortcut(QKeySequence("Ctrl++"))
        zoom_in_action.triggered.connect(self._zoom_in)
        zoom_out_action = zoom_menu.addAction("Zoom Out")
        zoom_out_action.setShortcut(QKeySequence("Ctrl+-"))
        zoom_out_action.triggered.connect(self._zoom_out)
        zoom_menu.addSeparator()
        fit_width_action = zoom_menu.addAction("Fit Width")
        fit_width_action.setShortcut(QKeySequence("Ctrl+1"))
        fit_width_action.triggered.connect(self._fit_width)
        fit_page_action = zoom_menu.addAction("Fit Page")
        fit_page_action.setShortcut(QKeySequence("Ctrl+2"))
        fit_page_action.triggered.connect(self._fit_page)
        zoom_100_action = zoom_menu.addAction("100%")
        zoom_100_action.setShortcut(QKeySequence("Ctrl+0"))
        zoom_100_action.triggered.connect(lambda: self._viewer.set_zoom(100))

        # Rotation submenu
        rotate_menu = view_menu.addMenu("Rotate")
        assert rotate_menu is not None
        rotate_cw_action = rotate_menu.addAction("Rotate Clockwise")
        rotate_cw_action.setShortcut(QKeySequence("Ctrl+R"))
        rotate_cw_action.triggered.connect(lambda: self._rotate(90))
        rotate_ccw_action = rotate_menu.addAction("Rotate Counter-Clockwise")
        rotate_ccw_action.setShortcut(QKeySequence("Ctrl+Shift+R"))
        rotate_ccw_action.triggered.connect(lambda: self._rotate(-90))

        view_menu.addSeparator()

        # View mode submenu
        view_mode_menu = view_menu.addMenu("View Mode")
        assert view_mode_menu is not None
        view_mode_menu.addAction(self._create_action(
            "Single Page", lambda: self._set_view_mode(ViewMode.SINGLE_PAGE)))
        view_mode_menu.addAction(self._create_action(
            "Two Pages", lambda: self._set_view_mode(ViewMode.TWO_PAGE)))
        view_mode_menu.addAction(self._create_action(
            "Continuous", lambda: self._set_view_mode(ViewMode.CONTINUOUS)))

        view_menu.addSeparator()

        # Sidebar toggle
        self._sidebar_action = view_menu.addAction("Show Sidebar")
        self._sidebar_action.setCheckable(True)
        self._sidebar_action.setChecked(True)
        self._sidebar_action.triggered.connect(self._toggle_sidebar)

        # Toolbar toggles
        self._main_toolbar_action = view_menu.addAction("Show Main Toolbar")
        self._main_toolbar_action.setCheckable(True)
        self._main_toolbar_action.setChecked(True)

        self._annotation_toolbar_action = view_menu.addAction(
            "Show Annotation Toolbar")
        self._annotation_toolbar_action.setCheckable(True)
        self._annotation_toolbar_action.setChecked(True)

        view_menu.addSeparator()

        fullscreen_action = self._create_action(
            "Full Screen", self._toggle_fullscreen, QKeySequence("F11"))
        view_menu.addAction(fullscreen_action)

        # === Page Menu ===
        page_menu = menubar.addMenu("&Page")
        assert page_menu is not None

        page_menu.addAction(self._create_action(
            "Insert Blank Page...", self._insert_blank_page))
        page_menu.addAction(self._create_action(
            "Insert Pages from File...", self._insert_from_file))
        page_menu.addSeparator()
        page_menu.addAction(self._create_action(
            "Delete Page", self._delete_current_page))
        page_menu.addAction(self._create_action(
            "Extract Pages...", self._extract_pages))
        page_menu.addSeparator()
        page_menu.addAction(self._create_action(
            "Rotate Clockwise", lambda: self._rotate_page(90)))
        page_menu.addAction(self._create_action(
            "Rotate Counter-Clockwise", lambda: self._rotate_page(-90)))
        page_menu.addSeparator()
        page_menu.addAction(self._create_action(
            "Crop Page...", self._crop_page))

        # === Tools Menu ===
        tools_menu = menubar.addMenu("&Tools")
        assert tools_menu is not None

        tools_menu.addAction(self._create_action(
            "Merge PDFs...", self._merge_pdfs))
        tools_menu.addAction(self._create_action(
            "Split PDF...", self._split_pdf))
        tools_menu.addSeparator()
        tools_menu.addAction(self._create_action(
            "Compress PDF...", self._compress_pdf))
        tools_menu.addAction(self._create_action(
            "Optimize PDF...", self._optimize_pdf))
        tools_menu.addSeparator()
        tools_menu.addAction(self._create_action(
            "OCR (Recognize Text)...", self._run_ocr))
        tools_menu.addSeparator()
        tools_menu.addAction(self._create_action(
            "Add Watermark...", self._add_watermark))
        tools_menu.addAction(self._create_action(
            "Add Header/Footer...", self._add_header_footer))
        tools_menu.addSeparator()
        tools_menu.addAction(self._create_action(
            "Encrypt PDF...", self._encrypt_pdf))
        tools_menu.addAction(self._create_action(
            "Remove Password...", self._remove_password))
        tools_menu.addSeparator()
        tools_menu.addAction(self._create_action(
            "Batch Process...", self._batch_process))

        # === Help Menu ===
        help_menu = menubar.addMenu("&Help")
        assert help_menu is not None
        help_menu.addAction(self._create_action("&About", self._show_about))
        help_menu.addAction(self._create_action(
            "Keyboard Shortcuts", self._show_shortcuts))

    def _setup_toolbars(self):
        """Setup toolbars"""
        # Main toolbar
        self._main_toolbar = MainToolbar()
        self.addToolBar(Qt.ToolBarArea.TopToolBarArea, self._main_toolbar)

        # Add toolbar break to put annotation toolbar on its own row
        self.addToolBarBreak(Qt.ToolBarArea.TopToolBarArea)

        # Annotation toolbar (on second row for better visibility)
        self._annotation_toolbar = AnnotationToolbar()
        self.addToolBar(Qt.ToolBarArea.TopToolBarArea,
                        self._annotation_toolbar)

        # Connect toolbar signals
        self._main_toolbar.open_requested.connect(self._open_document)
        self._main_toolbar.save_requested.connect(self._save_document)
        self._main_toolbar.save_as_requested.connect(self._save_document_as)
        self._main_toolbar.print_requested.connect(self._print_document)
        self._main_toolbar.undo_requested.connect(self._undo)
        self._main_toolbar.redo_requested.connect(self._redo)
        self._main_toolbar.zoom_changed.connect(self._viewer.set_zoom)
        self._main_toolbar.zoom_in_requested.connect(self._zoom_in)
        self._main_toolbar.zoom_out_requested.connect(self._zoom_out)
        self._main_toolbar.fit_width_requested.connect(self._fit_width)
        self._main_toolbar.fit_page_requested.connect(self._fit_page)
        self._main_toolbar.rotate_cw_requested.connect(
            lambda: self._rotate(90))
        self._main_toolbar.rotate_ccw_requested.connect(
            lambda: self._rotate(-90))
        self._main_toolbar.page_changed.connect(self._viewer.go_to_page)
        self._main_toolbar.first_page_requested.connect(
            self._viewer.first_page)
        self._main_toolbar.prev_page_requested.connect(
            self._viewer.previous_page)
        self._main_toolbar.next_page_requested.connect(self._viewer.next_page)
        self._main_toolbar.last_page_requested.connect(self._viewer.last_page)
        self._main_toolbar.search_requested.connect(self._search)

        # Annotation toolbar signals
        self._annotation_toolbar.tool_changed.connect(self._on_tool_changed)
        self._annotation_toolbar.color_changed.connect(
            self._viewer.set_annotation_color)
        self._annotation_toolbar.opacity_changed.connect(
            self._viewer.set_annotation_opacity)
        self._annotation_toolbar.stroke_width_changed.connect(
            self._viewer.set_stroke_width)
        self._annotation_toolbar.font_changed.connect(self._viewer.set_font)

    def _setup_statusbar(self):
        """Setup status bar"""
        self._statusbar = QStatusBar()
        self.setStatusBar(self._statusbar)

        # Page info label
        self._page_info_label = QLabel("No document")
        self._statusbar.addWidget(self._page_info_label)

        # Spacer
        self._statusbar.addWidget(QWidget(), 1)

        # Zoom info label
        self._zoom_label = QLabel("100%")
        self._statusbar.addPermanentWidget(self._zoom_label)

        # File size label
        self._file_size_label = QLabel("")
        self._statusbar.addPermanentWidget(self._file_size_label)

    def _connect_signals(self):
        """Connect signals between components"""
        # Viewer signals
        self._viewer.page_changed.connect(self._on_page_changed)
        self._viewer.zoom_changed.connect(self._on_zoom_changed)
        self._viewer.document_modified.connect(self._on_document_modified)
        self._viewer.annotation_create_requested.connect(
            self._create_annotation)

        # Sidebar signals
        self._sidebar.page_selected.connect(self._viewer.go_to_page)
        self._sidebar.page_double_clicked.connect(self._viewer.go_to_page)
        self._sidebar.bookmark_clicked.connect(self._viewer.go_to_page)
        self._sidebar.page_rotate_requested.connect(self._rotate_page)
        self._sidebar.page_delete_requested.connect(self._delete_page)
        self._sidebar.page_extract_requested.connect(
            self._extract_specific_pages)

    def _apply_settings(self):
        """Apply saved settings"""
        # Restore window geometry if available
        if self._settings.window_geometry:
            try:
                import base64
                geometry = base64.b64decode(self._settings.window_geometry)
                self.restoreGeometry(geometry)
            except Exception:
                pass

        # Restore sidebar visibility
        self._sidebar.setVisible(self._settings.sidebar_visible)
        self._sidebar_action.setChecked(self._settings.sidebar_visible)

    def _save_settings(self):
        """Save current settings"""
        import base64
        geometry_data = self.saveGeometry().data()  # Get raw bytes from QByteArray
        self._settings.window_geometry = base64.b64encode(geometry_data).decode('utf-8')  # type: ignore[assignment]
        self._settings.sidebar_visible = self._sidebar.isVisible()
        self._settings.save(config.SETTINGS_PATH)

    # ==================== File Operations ====================

    def _new_document(self):
        """Create a new document"""
        if not self._confirm_close():
            return

        try:
            self._document.create_new()
            self._document.add_blank_page()
            self._current_file = None
            self._load_document_to_viewer()
            self._is_modified = True
            self._update_title()
        except Exception as e:
            QMessageBox.critical(
                self, "Error", f"Failed to create new document:\n{e}")

    def _open_document(self):
        """Open a document"""
        if not self._confirm_close():
            return

        filepath, _ = QFileDialog.getOpenFileName(
            self,
            "Open PDF",
            self._settings.last_opened_directory,
            "PDF Files (*.pdf);;All Files (*)"
        )

        if filepath:
            self._open_file(filepath)

    def _open_file(self, filepath: str):
        """Open a specific file"""
        try:
            path = Path(filepath)
            self._settings.last_opened_directory = str(path.parent)

            # Check if password needed
            result = self._document.open(path)

            if not result and self._document.needs_password:
                # Ask for password
                from PyQt6.QtWidgets import QLineEdit
                password, ok = QInputDialog.getText(
                    self, "Password Required",
                    "This PDF is password protected. Enter password:",
                    echo=QLineEdit.EchoMode.Password
                )
                if ok:
                    self._document.open(path, password)
                else:
                    return

            self._load_document_to_viewer()
            self._current_file = path
            self._is_modified = False
            self._update_title()

            # Add to recent files
            self._settings.add_recent_file(str(path))
            self._update_recent_files_menu()

            # Update file size
            size = path.stat().st_size
            self._file_size_label.setText(self._format_size(size))

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to open file:\n{e}")

    def _load_document_to_viewer(self):
        """Load the current document into the viewer"""
        if self._document.is_open:
            self._viewer.set_document(self._document._doc, str(
                self._current_file) if self._current_file else None)
            self._sidebar.set_document(self._document._doc)
            self._main_toolbar.set_page_count(self._document.page_count)
            self._update_actions_state()

    def _save_document(self):
        """Save the current document"""
        if not self._document.is_open:
            return

        if self._current_file:
            try:
                self._document.save()
                self._is_modified = False
                self._update_title()
                self._statusbar.showMessage("Document saved", 3000)
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to save:\n{e}")
        else:
            self._save_document_as()

    def _save_document_as(self):
        """Save document with new name"""
        if not self._document.is_open:
            return

        filepath, _ = QFileDialog.getSaveFileName(
            self,
            "Save PDF As",
            self._settings.last_saved_directory,
            "PDF Files (*.pdf)"
        )

        if filepath:
            try:
                path = Path(filepath)
                if not path.suffix.lower() == '.pdf':
                    path = path.with_suffix('.pdf')

                self._settings.last_saved_directory = str(path.parent)
                self._document.save(path)
                self._current_file = path
                self._is_modified = False
                self._update_title()
                self._statusbar.showMessage("Document saved", 3000)
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to save:\n{e}")

    def _close_document(self):
        """Close the current document"""
        if not self._confirm_close():
            return

        self._document.close()
        self._viewer.set_document(None, None)
        self._sidebar.set_document(None)
        self._current_file = None
        self._is_modified = False
        self._update_title()
        self._update_actions_state()

    def _confirm_close(self) -> bool:
        """Confirm closing with unsaved changes"""
        if not self._is_modified:
            return True

        result = QMessageBox.question(
            self,
            "Unsaved Changes",
            "Do you want to save changes before closing?",
            QMessageBox.StandardButton.Save |
            QMessageBox.StandardButton.Discard |
            QMessageBox.StandardButton.Cancel
        )

        if result == QMessageBox.StandardButton.Save:
            self._save_document()
            return not self._is_modified
        elif result == QMessageBox.StandardButton.Discard:
            return True
        else:
            return False

    def _print_document(self):
        """Print the current document"""
        if not self._document.is_open:
            return

        try:
            from PyQt6.QtPrintSupport import QPrinter, QPrintDialog
            from PyQt6.QtGui import QPainter
        except ImportError:
            QMessageBox.warning(self, "Print Not Available",
                                "Print support requires PyQt6 print modules.")
            return

        printer = QPrinter(QPrinter.PrinterMode.HighResolution)
        printer.setPageOrientation(QPageLayout.Orientation.Portrait)

        dialog = QPrintDialog(printer, self)
        dialog.setWindowTitle("Print Document")

        if dialog.exec() == QPrintDialog.DialogCode.Accepted:
            # Create progress dialog
            progress = QProgressDialog(
                "Printing...", "Cancel", 0, self._document.page_count, self)
            progress.setWindowTitle("Printing")
            progress.setWindowModality(Qt.WindowModality.WindowModal)
            progress.show()

            painter = QPainter()
            painter.begin(printer)

            try:
                for i in range(self._document.page_count):
                    if progress.wasCanceled():
                        break

                    progress.setValue(i)
                    progress.setLabelText(
                        f"Printing page {i + 1} of {self._document.page_count}...")
                    QApplication.processEvents()

                    if i > 0:
                        printer.newPage()

                    # Render page to image at printer resolution
                    page = self._document._doc[i]
                    # Calculate scale for printer DPI
                    dpi = printer.resolution()
                    scale = dpi / 72.0  # PDF points to printer DPI
                    mat = fitz.Matrix(scale, scale)
                    pix = page.get_pixmap(matrix=mat)

                    # Convert to QImage
                    img = QImage(pix.samples, pix.width, pix.height,
                                 pix.stride, QImage.Format.Format_RGB888)
                    pixmap = QPixmap.fromImage(img)

                    # Calculate position to center on page
                    page_rect = printer.pageRect(QPrinter.Unit.DevicePixel)
                    x = (page_rect.width() - pixmap.width()) / 2
                    y = (page_rect.height() - pixmap.height()) / 2

                    # Scale to fit page if needed
                    if pixmap.width() > page_rect.width() or pixmap.height() > page_rect.height():
                        pixmap = pixmap.scaled(
                            int(page_rect.width()), int(page_rect.height()),
                            Qt.AspectRatioMode.KeepAspectRatio,
                            Qt.TransformationMode.SmoothTransformation
                        )
                        x = (page_rect.width() - pixmap.width()) / 2
                        y = (page_rect.height() - pixmap.height()) / 2

                    painter.drawPixmap(int(x), int(y), pixmap)

                progress.setValue(self._document.page_count)
                self._statusbar.showMessage("Document sent to printer", 3000)

            except Exception as e:
                QMessageBox.critical(self, "Print Error",
                                     f"Failed to print:\n{e}")
            finally:
                painter.end()

    def _show_properties(self):
        """Show document properties"""
        if not self._document.is_open:
            return

        metadata = self._document.get_metadata()
        info = f"""
Title: {metadata.title}
Author: {metadata.author}
Subject: {metadata.subject}
Keywords: {metadata.keywords}
Creator: {metadata.creator}
Producer: {metadata.producer}
Created: {metadata.creation_date}
Modified: {metadata.modification_date}
Pages: {metadata.page_count}
Size: {self._format_size(metadata.file_size)}
Encrypted: {metadata.encryption}
        """
        QMessageBox.information(self, "Document Properties", info.strip())

    # ==================== Edit Operations ====================

    def _undo(self):
        """Undo last action"""
        if self._history_manager.can_undo():
            desc = self._history_manager.get_undo_description()
            if self._history_manager.undo():
                self._load_document_to_viewer()
                self._is_modified = True
                self._update_title()
                self._statusbar.showMessage(f"Undo: {desc}", 2000)
        else:
            self._statusbar.showMessage("Nothing to undo", 2000)

    def _redo(self):
        """Redo last undone action"""
        if self._history_manager.can_redo():
            desc = self._history_manager.get_redo_description()
            if self._history_manager.redo():
                self._load_document_to_viewer()
                self._is_modified = True
                self._update_title()
                self._statusbar.showMessage(f"Redo: {desc}", 2000)
        else:
            self._statusbar.showMessage("Nothing to redo", 2000)

    def _cut(self):
        """Cut selection"""
        self._copy()
        self._delete()

    def _copy(self):
        """Copy selection"""
        # Copy selected text to clipboard
        selected_text = self._viewer.get_selected_text()
        if selected_text:
            clipboard = QApplication.clipboard()
            clipboard.setText(selected_text)
            self._statusbar.showMessage("Text copied to clipboard", 2000)
        else:
            self._statusbar.showMessage("No text selected", 2000)

    def _paste(self):
        """Paste from clipboard"""
        clipboard = QApplication.clipboard()
        text = clipboard.text()
        if text:
            # For now, show info that paste creates text annotation
            self._statusbar.showMessage(
                "Use Text Box tool to add text to PDF", 3000)
        else:
            self._statusbar.showMessage("Clipboard is empty", 2000)

    def _delete(self):
        """Delete selection"""
        # Delete selected annotation if any
        self._statusbar.showMessage("Select an annotation to delete", 2000)

    def _select_all(self):
        """Select all text on current page"""
        if not self._document.is_open:
            return
        page_num = self._viewer.get_current_page()
        text = self._document.get_page_text(page_num)
        if text:
            clipboard = QApplication.clipboard()
            clipboard.setText(text)
            self._statusbar.showMessage("Page text copied to clipboard", 2000)

    def _show_find(self):
        """Show find dialog"""
        if not self._document.is_open:
            return

        if self._find_dialog is None:
            self._find_dialog = FindDialog(self)
            self._find_dialog.find_requested.connect(self._do_search)
            self._find_dialog.find_next.connect(self._find_next)
            self._find_dialog.find_prev.connect(self._find_previous)
            self._find_dialog.closed.connect(self._on_find_closed)

        self._find_dialog.show()
        self._find_dialog.raise_()
        self._find_dialog.focus_search()

    def _show_replace(self):
        """Show find and replace dialog"""
        if not self._document.is_open:
            return

        if self._replace_dialog is None:
            self._replace_dialog = FindReplaceDialog(self)
            self._replace_dialog.find_requested.connect(self._do_search)
            self._replace_dialog.find_next.connect(self._find_next)
            self._replace_dialog.find_prev.connect(self._find_previous)
            self._replace_dialog.replace_requested.connect(self._do_replace)
            self._replace_dialog.replace_all_requested.connect(
                self._do_replace_all)
            self._replace_dialog.closed.connect(self._on_replace_closed)

        self._replace_dialog.show()
        self._replace_dialog.raise_()
        self._replace_dialog.focus_search()

    def _do_search(self, text: str, case_sensitive: bool = False):
        """Perform search and update results"""
        if not self._document.is_open or not text:
            self._search_results = []
            return

        self._search_results = self._document.search_text(text, case_sensitive)
        total = sum(len(r['rects']) for r in self._search_results)

        if self._search_results:
            self._current_search_index = 0
            self._go_to_search_result(0)

            if self._find_dialog:
                self._find_dialog.set_result_count(0, total)
            if self._replace_dialog:
                self._replace_dialog.set_result_count(0, total)
        else:
            if self._find_dialog:
                self._find_dialog.set_result_count(0, 0)
            if self._replace_dialog:
                self._replace_dialog.set_result_count(0, 0)

    def _find_next(self):
        """Go to next search result"""
        if not self._search_results:
            return

        total = sum(len(r['rects']) for r in self._search_results)
        self._current_search_index = (self._current_search_index + 1) % total
        self._go_to_search_result(self._current_search_index)

        if self._find_dialog:
            self._find_dialog.set_result_count(
                self._current_search_index, total)
        if self._replace_dialog:
            self._replace_dialog.set_result_count(
                self._current_search_index, total)

    def _find_previous(self):
        """Go to previous search result"""
        if not self._search_results:
            return

        total = sum(len(r['rects']) for r in self._search_results)
        self._current_search_index = (self._current_search_index - 1) % total
        self._go_to_search_result(self._current_search_index)

        if self._find_dialog:
            self._find_dialog.set_result_count(
                self._current_search_index, total)
        if self._replace_dialog:
            self._replace_dialog.set_result_count(
                self._current_search_index, total)

    def _go_to_search_result(self, index: int):
        """Navigate to a specific search result"""
        current = 0
        for result in self._search_results:
            page_num = result['page']
            for rect in result['rects']:
                if current == index:
                    self._viewer.go_to_page(page_num)
                    return
                current += 1

    def _do_replace(self, replacement: str):
        """Replace current occurrence"""
        # PDF text replacement is complex - using redaction approach
        QMessageBox.information(
            self,
            "Replace",
            "Direct text replacement in PDFs is limited.\n"
            "Consider using the Redaction tool to remove text and then add new text."
        )

    def _do_replace_all(self, replacement: str):
        """Replace all occurrences"""
        QMessageBox.information(
            self,
            "Replace All",
            "Direct text replacement in PDFs is limited.\n"
            "PDF files store text as rendered graphics, not editable text.\n"
            "Consider exporting to Word format for text editing."
        )

    def _on_find_closed(self):
        """Handle find dialog closed"""
        self._search_results = []

    def _on_replace_closed(self):
        """Handle replace dialog closed"""
        self._search_results = []

    def _search(self, text: str):
        """Search for text (from toolbar)"""
        if not self._document.is_open or not text:
            return

        self._do_search(text, False)
        total = sum(len(r['rects']) for r in self._search_results)
        if total > 0:
            self._statusbar.showMessage(f"Found {total} matches")
        else:
            self._statusbar.showMessage("No matches found")

    # ==================== View Operations ====================

    def _zoom_in(self):
        """Zoom in"""
        self._viewer.zoom_in()

    def _zoom_out(self):
        """Zoom out"""
        self._viewer.zoom_out()

    def _fit_width(self):
        """Fit to width"""
        self._viewer.fit_width()

    def _fit_page(self):
        """Fit whole page"""
        self._viewer.fit_page()

    def _rotate(self, degrees: int):
        """Rotate view"""
        self._viewer.rotate_view(degrees)

    def _set_view_mode(self, mode: ViewMode):
        """Set view mode"""
        self._viewer.set_view_mode(mode)
        self._statusbar.showMessage(f"View mode: {mode.value}", 2000)

    def _toggle_sidebar(self, visible: bool):
        """Toggle sidebar visibility"""
        self._sidebar.setVisible(visible)

    def _toggle_fullscreen(self):
        """Toggle fullscreen mode"""
        if self.isFullScreen():
            self.showNormal()
        else:
            self.showFullScreen()

    # ==================== Page Operations ====================

    def _insert_blank_page(self):
        """Insert a blank page"""
        if not self._document.is_open:
            return

        current = self._viewer.get_current_page()
        command = PageAddCommand(self._document, current + 1)
        if self._history_manager.execute(command):
            self._load_document_to_viewer()
            self._is_modified = True
            self._update_title()
            self._statusbar.showMessage("Inserted blank page", 2000)

    def _insert_from_file(self):
        """Insert pages from another PDF"""
        if not self._document.is_open:
            return

        filepath, _ = QFileDialog.getOpenFileName(
            self, "Select PDF to Insert", "", "PDF Files (*.pdf)"
        )
        if filepath:
            current = self._viewer.get_current_page()
            self._document.merge_pdf(filepath, current + 1)
            self._load_document_to_viewer()
            self._is_modified = True
            self._update_title()

    def _delete_current_page(self):
        """Delete the current page"""
        self._delete_page(self._viewer.get_current_page())

    def _delete_page(self, page_num: int):
        """Delete a specific page"""
        if not self._document.is_open:
            return

        if self._document.page_count <= 1:
            QMessageBox.warning(self, "Cannot Delete",
                                "Cannot delete the only page in the document.")
            return

        result = QMessageBox.question(
            self, "Delete Page",
            f"Are you sure you want to delete page {page_num + 1}?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )

        if result == QMessageBox.StandardButton.Yes:
            command = PageDeleteCommand(self._document, page_num)
            if self._history_manager.execute(command):
                self._load_document_to_viewer()
                self._is_modified = True
                self._update_title()
                self._statusbar.showMessage("Deleted page", 2000)

    def _rotate_page(self, page_num: int, degrees: int = 90):
        """Rotate a specific page"""
        if not self._document.is_open:
            return

        command = PageRotateCommand(self._document, page_num, degrees)
        if self._history_manager.execute(command):
            self._load_document_to_viewer()
            self._is_modified = True
            self._update_title()
            self._statusbar.showMessage("Rotated page", 2000)

    def _extract_pages(self):
        """Extract pages to new PDF"""
        if not self._document.is_open:
            return

        dialog = ExtractPagesDialog(
            self._document.page_count, self._viewer.get_current_page(), self)
        if dialog.exec():
            pages = dialog.get_selected_pages()
            if pages:
                filepath, _ = QFileDialog.getSaveFileName(
                    self, "Save Extracted Pages", "", "PDF Files (*.pdf)"
                )
                if filepath:
                    try:
                        self._document.extract_pages(pages, filepath)
                        self._statusbar.showMessage(
                            f"Extracted {len(pages)} pages", 3000)

                        if QMessageBox.question(
                            self, "Open Extracted PDF",
                            "Do you want to open the extracted PDF?",
                            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
                        ) == QMessageBox.StandardButton.Yes:
                            self._open_file(filepath)
                    except Exception as e:
                        QMessageBox.critical(
                            self, "Error", f"Failed to extract pages:\n{e}")

    def _extract_specific_pages(self, pages: List[int]):
        """Extract specific pages"""
        if not self._document.is_open or not pages:
            return

        filepath, _ = QFileDialog.getSaveFileName(
            self, "Save Extracted Pages", "", "PDF Files (*.pdf)"
        )
        if filepath:
            self._document.extract_pages(pages, filepath)
            self._statusbar.showMessage(f"Extracted {len(pages)} pages", 3000)

    def _crop_page(self):
        """Crop current page"""
        if not self._document.is_open:
            return

        page_num = self._viewer.get_current_page()
        page = self._document._doc[page_num]
        rect = page.rect

        # Render page for preview - create a QPixmap
        pix = page.get_pixmap(matrix=fitz.Matrix(0.5, 0.5))
        img = QImage(pix.samples, pix.width, pix.height,
                     pix.stride, QImage.Format.Format_RGB888)
        preview_pixmap = QPixmap.fromImage(img)

        dialog = CropDialog(preview_pixmap, rect.width, rect.height, self)
        if dialog.exec():
            crop_rect = dialog.get_crop_rect()
            try:
                # Apply crop to the page
                new_rect = fitz.Rect(
                    crop_rect[0], crop_rect[1], crop_rect[2], crop_rect[3])
                page.set_cropbox(new_rect)

                # Apply to all pages if requested
                if dialog.apply_to_all_pages():
                    for i in range(self._document.page_count):
                        if i != page_num:
                            p = self._document._doc[i]
                            p_rect = p.rect
                            # Scale crop proportionally
                            scale_x = p_rect.width / rect.width
                            scale_y = p_rect.height / rect.height
                            p_new_rect = fitz.Rect(
                                crop_rect[0] * scale_x,
                                crop_rect[1] * scale_y,
                                crop_rect[2] * scale_x,
                                crop_rect[3] * scale_y
                            )
                            p.set_cropbox(p_new_rect)

                self._load_document_to_viewer()
                self._is_modified = True
                self._update_title()
                self._statusbar.showMessage("Page(s) cropped", 2000)
            except Exception as e:
                QMessageBox.critical(
                    self, "Error", f"Failed to crop page:\n{e}")

    # ==================== Tools ====================

    def _merge_pdfs(self):
        """Merge multiple PDFs"""
        files, _ = QFileDialog.getOpenFileNames(
            self, "Select PDFs to Merge", "", "PDF Files (*.pdf)"
        )
        if len(files) < 2:
            return

        output, _ = QFileDialog.getSaveFileName(
            self, "Save Merged PDF", "", "PDF Files (*.pdf)"
        )
        if output:
            try:
                self._document.merge_pdfs(files, output)
                self._statusbar.showMessage("PDFs merged successfully", 3000)

                if QMessageBox.question(
                    self, "Open Merged PDF",
                    "Do you want to open the merged PDF?",
                    QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
                ) == QMessageBox.StandardButton.Yes:
                    self._open_file(output)
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to merge:\n{e}")

    def _split_pdf(self):
        """Split PDF into multiple files"""
        if not self._document.is_open:
            return

        output_dir = QFileDialog.getExistingDirectory(
            self, "Select Output Directory")
        if output_dir:
            try:
                files = self._document.split_by_pages(output_dir, 1)
                self._statusbar.showMessage(
                    f"Split into {len(files)} files", 3000)
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to split:\n{e}")

    def _compress_pdf(self):
        """Compress PDF to reduce size"""
        if not self._document.is_open:
            return

        filepath, _ = QFileDialog.getSaveFileName(
            self, "Save Compressed PDF", "", "PDF Files (*.pdf)"
        )
        if filepath:
            try:
                self._document.compress(filepath)
                self._statusbar.showMessage("PDF compressed", 3000)
            except Exception as e:
                QMessageBox.critical(
                    self, "Error", f"Failed to compress:\n{e}")

    def _optimize_pdf(self):
        """Optimize PDF"""
        self._compress_pdf()

    def _run_ocr(self):
        """Run OCR on document"""
        if not self._document.is_open:
            return

        try:
            import pytesseract
            from PIL import Image
            import io
        except ImportError:
            QMessageBox.warning(
                self,
                "OCR Not Available",
                "OCR requires pytesseract and Pillow libraries.\n\n"
                "Install with: pip install pytesseract Pillow\n\n"
                "You also need Tesseract OCR installed on your system:\n"
                "- Windows: Download from https://github.com/UB-Mannheim/tesseract/wiki\n"
                "- macOS: brew install tesseract\n"
                "- Linux: sudo apt install tesseract-ocr"
            )
            return

        # Confirm with user
        result = QMessageBox.question(
            self,
            "Run OCR",
            f"This will add a searchable text layer to all {self._document.page_count} pages.\n\n"
            "This process may take some time. Continue?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )
        if result != QMessageBox.StandardButton.Yes:
            return

        # Create progress dialog
        progress = QProgressDialog(
            "Running OCR...", "Cancel", 0, self._document.page_count, self)
        progress.setWindowTitle("OCR Processing")
        progress.setWindowModality(Qt.WindowModality.WindowModal)
        progress.show()

        try:
            for i in range(self._document.page_count):
                if progress.wasCanceled():
                    break

                progress.setValue(i)
                progress.setLabelText(
                    f"Processing page {i + 1} of {self._document.page_count}...")
                QApplication.processEvents()

                # Render page to image
                page = self._document._doc[i]
                # Higher resolution for better OCR
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                img_data = pix.tobytes("png")

                # OCR the image
                img = Image.open(io.BytesIO(img_data))
                text = pytesseract.image_to_string(img)

                # Add text layer to page (invisible)
                if text.strip():
                    # Insert as invisible text behind the image
                    text_point = fitz.Point(0, page.rect.height)
                    page.insert_text(text_point, text, fontsize=1,
                                     color=(1, 1, 1), render_mode=3)

            progress.setValue(self._document.page_count)
            self._load_document_to_viewer()
            self._is_modified = True
            self._update_title()
            self._statusbar.showMessage(
                "OCR completed - document is now searchable", 3000)

        except Exception as e:
            QMessageBox.critical(self, "OCR Error", f"Failed to run OCR:\n{e}")

    def _add_watermark(self):
        """Add watermark to document"""
        if not self._document.is_open:
            return

        text, ok = QInputDialog.getText(
            self, "Add Watermark", "Watermark text:")
        if ok and text:
            self._document.add_watermark(text)
            self._load_document_to_viewer()
            self._is_modified = True
            self._update_title()

    def _add_header_footer(self):
        """Add header/footer"""
        if not self._document.is_open:
            return

        dialog = HeaderFooterDialog(self._document.page_count, self)
        if dialog.exec():
            try:
                # Get settings from dialog
                header_texts = dialog.get_header_texts()
                footer_texts = dialog.get_footer_texts()
                font_settings = dialog.get_font_settings()
                page_range = dialog.get_page_range()
                margins = dialog.get_margins()

                fontsize = font_settings['size']
                margin_top = margins['top']
                margin_bottom = margins['bottom']
                margin_side = margins['side']

                for i in range(page_range[0], page_range[1] + 1):
                    page = self._document._doc[i]
                    rect = page.rect

                    # Process variables in text
                    def process_text(text: str) -> str:
                        text = text.replace("{page}", str(i + 1))
                        text = text.replace("{total}", str(
                            self._document.page_count))
                        text = text.replace(
                            "{date}", datetime.datetime.now().strftime("%Y-%m-%d"))
                        return text

                    # Add header
                    if header_texts['left'] or header_texts['center'] or header_texts['right']:
                        y_pos = margin_top

                        if header_texts['left']:
                            text = process_text(header_texts['left'])
                            page.insert_text(
                                fitz.Point(margin_side, y_pos),
                                text, fontsize=fontsize
                            )
                        if header_texts['center']:
                            text = process_text(header_texts['center'])
                            text_width = fitz.get_text_length(
                                text, fontsize=fontsize)
                            x_pos = (rect.width - text_width) / 2
                            page.insert_text(fitz.Point(
                                x_pos, y_pos), text, fontsize=fontsize)
                        if header_texts['right']:
                            text = process_text(header_texts['right'])
                            text_width = fitz.get_text_length(
                                text, fontsize=fontsize)
                            x_pos = rect.width - margin_side - text_width
                            page.insert_text(fitz.Point(
                                x_pos, y_pos), text, fontsize=fontsize)

                    # Add footer
                    if footer_texts['left'] or footer_texts['center'] or footer_texts['right']:
                        y_pos = rect.height - margin_bottom

                        if footer_texts['left']:
                            text = process_text(footer_texts['left'])
                            page.insert_text(
                                fitz.Point(margin_side, y_pos),
                                text, fontsize=fontsize
                            )
                        if footer_texts['center']:
                            text = process_text(footer_texts['center'])
                            text_width = fitz.get_text_length(
                                text, fontsize=fontsize)
                            x_pos = (rect.width - text_width) / 2
                            page.insert_text(fitz.Point(
                                x_pos, y_pos), text, fontsize=fontsize)
                        if footer_texts['right']:
                            text = process_text(footer_texts['right'])
                            text_width = fitz.get_text_length(
                                text, fontsize=fontsize)
                            x_pos = rect.width - margin_side - text_width
                            page.insert_text(fitz.Point(
                                x_pos, y_pos), text, fontsize=fontsize)

                self._load_document_to_viewer()
                self._is_modified = True
                self._update_title()
                self._statusbar.showMessage("Header/Footer added", 2000)

            except Exception as e:
                QMessageBox.critical(
                    self, "Error", f"Failed to add header/footer:\n{e}")

    def _encrypt_pdf(self):
        """Encrypt PDF with password"""
        if not self._document.is_open:
            return

        from PyQt6.QtWidgets import QLineEdit
        password, ok = QInputDialog.getText(
            self, "Encrypt PDF", "Enter password:",
            QLineEdit.EchoMode.Password
        )
        if ok and password:
            self._document.encrypt(
                user_password=password, owner_password=password)
            self._save_document()

    def _remove_password(self):
        """Remove password protection"""
        if not self._document.is_open:
            return

        if not self._document._doc.is_encrypted:
            QMessageBox.information(
                self, "No Password", "This document is not password protected.")
            return

        # Ask for current password
        from PyQt6.QtWidgets import QLineEdit
        password, ok = QInputDialog.getText(
            self, "Remove Password",
            "Enter current password to remove protection:",
            QLineEdit.EchoMode.Password
        )
        if not ok:
            return

        try:
            # Try to authenticate with the password
            if self._document._doc.authenticate(password):
                # Save without encryption
                filepath, _ = QFileDialog.getSaveFileName(
                    self, "Save Unprotected PDF", "", "PDF Files (*.pdf)"
                )
                if filepath:
                    self._document._doc.save(
                        filepath, encryption=0)  # PDF_ENCRYPT_NONE
                    self._statusbar.showMessage(
                        "Password removed and saved", 3000)

                    if QMessageBox.question(
                        self, "Open Unprotected PDF",
                        "Do you want to open the unprotected PDF?",
                        QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
                    ) == QMessageBox.StandardButton.Yes:
                        self._open_file(filepath)
            else:
                QMessageBox.warning(self, "Incorrect Password",
                                    "The password you entered is incorrect.")
        except Exception as e:
            QMessageBox.critical(
                self, "Error", f"Failed to remove password:\n{e}")

    def _batch_process(self):
        """Open batch processing dialog"""
        dialog = BatchDialog(self)
        dialog.exec()

    # ==================== Export ====================

    def _export_as_images(self):
        """Export pages as images"""
        if not self._document.is_open:
            return

        output_dir = QFileDialog.getExistingDirectory(
            self, "Select Output Directory")
        if output_dir:
            try:
                for i in range(self._document.page_count):
                    img_bytes = self._document.render_page_to_image(i)
                    filepath = Path(output_dir) / f"page_{i+1:04d}.png"
                    with open(filepath, 'wb') as f:
                        f.write(img_bytes)
                self._statusbar.showMessage(
                    f"Exported {self._document.page_count} images", 3000)
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to export:\n{e}")

    def _export_as_word(self):
        """Export as Word document"""
        if not self._document.is_open:
            return

        try:
            from docx import Document as WordDocument
            from docx.shared import Inches, Pt  # noqa: F401
        except ImportError:
            QMessageBox.warning(
                self,
                "Export Not Available",
                "Word export requires the python-docx library.\n\n"
                "Install with: pip install python-docx"
            )
            return

        filepath, _ = QFileDialog.getSaveFileName(
            self, "Export as Word Document", "", "Word Documents (*.docx)"
        )
        if not filepath:
            return

        # Create progress dialog
        progress = QProgressDialog(
            "Exporting to Word...", "Cancel", 0, self._document.page_count, self)
        progress.setWindowTitle("Exporting")
        progress.setWindowModality(Qt.WindowModality.WindowModal)
        progress.show()

        try:
            doc = WordDocument()

            for i in range(self._document.page_count):
                if progress.wasCanceled():
                    break

                progress.setValue(i)
                progress.setLabelText(
                    f"Processing page {i + 1} of {self._document.page_count}...")
                QApplication.processEvents()

                # Get text from page
                text = self._document.get_page_text(i)

                if text.strip():
                    # Add text paragraphs
                    for line in text.split('\n'):
                        if line.strip():
                            doc.add_paragraph(line)

                # Add page break between pages (except last page)
                if i < self._document.page_count - 1:
                    doc.add_page_break()

            progress.setValue(self._document.page_count)
            doc.save(filepath)
            self._statusbar.showMessage(
                f"Exported to {Path(filepath).name}", 3000)

            if QMessageBox.question(
                self, "Open Word Document",
                "Export complete. Do you want to open the Word document?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            ) == QMessageBox.StandardButton.Yes:
                import subprocess
                subprocess.Popen(['start', '', filepath], shell=True)

        except Exception as e:
            QMessageBox.critical(self, "Export Error",
                                 f"Failed to export:\n{e}")

    def _export_as_text(self):
        """Export as plain text"""
        if not self._document.is_open:
            return

        filepath, _ = QFileDialog.getSaveFileName(
            self, "Export as Text", "", "Text Files (*.txt)"
        )
        if filepath:
            try:
                text = self._document.get_all_text()
                with open(filepath, 'w', encoding='utf-8') as f:
                    f.write(text)
                self._statusbar.showMessage("Exported as text", 3000)
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to export:\n{e}")

    # ==================== Help ====================

    def _show_about(self):
        """Show about dialog"""
        QMessageBox.about(
            self,
            f"About {config.APP_NAME}",
            f"""<h2>{config.APP_NAME}</h2>
            <p>Version {config.APP_VERSION}</p>
            <p>A powerful yet simple PDF editor with all the features you need.</p>
            <p>Built with Python, PyQt6, and PyMuPDF.</p>
            """
        )

    def _show_shortcuts(self):
        """Show keyboard shortcuts"""
        shortcuts = """
<h3>Keyboard Shortcuts</h3>
<table>
<tr><td>Ctrl+N</td><td>New Document</td></tr>
<tr><td>Ctrl+O</td><td>Open</td></tr>
<tr><td>Ctrl+S</td><td>Save</td></tr>
<tr><td>Ctrl+Shift+S</td><td>Save As</td></tr>
<tr><td>Ctrl+P</td><td>Print</td></tr>
<tr><td>Ctrl+Z</td><td>Undo</td></tr>
<tr><td>Ctrl+Y</td><td>Redo</td></tr>
<tr><td>Ctrl+F</td><td>Find</td></tr>
<tr><td>Ctrl++</td><td>Zoom In</td></tr>
<tr><td>Ctrl+-</td><td>Zoom Out</td></tr>
<tr><td>Ctrl+0</td><td>Zoom 100%</td></tr>
<tr><td>Page Up/Down</td><td>Navigate Pages</td></tr>
<tr><td>Ctrl+Home</td><td>First Page</td></tr>
<tr><td>Ctrl+End</td><td>Last Page</td></tr>
<tr><td>F11</td><td>Full Screen</td></tr>
</table>
        """
        QMessageBox.information(self, "Keyboard Shortcuts", shortcuts)

    # ==================== Event Handlers ====================

    def _on_page_changed(self, page: int):
        """Handle page change"""
        self._main_toolbar.set_current_page(page)
        self._sidebar.set_current_page(page)
        self._page_info_label.setText(
            f"Page {page + 1} of {self._document.page_count}")

    def _on_zoom_changed(self, zoom: float):
        """Handle zoom change"""
        self._main_toolbar.set_zoom(zoom)
        self._zoom_label.setText(f"{zoom:.0f}%")

    def _on_document_modified(self):
        """Handle document modification"""
        self._is_modified = True
        self._update_title()
        # Refresh sidebar thumbnails to reflect changes
        self._sidebar.refresh()

    def _on_tool_changed(self, tool: str):
        """Handle tool change"""
        try:
            mode = ToolMode(tool)
            self._viewer.set_tool_mode(mode)
        except ValueError:
            pass  # Ignore unknown tool modes

    def _create_annotation(self, page_num: int, annot_type: str, rect: Any, data: Dict):
        """Handle annotation creation request"""
        if not self._document.is_open:
            return

        command = AnnotationAddCommand(
            self._document, page_num, annot_type, rect, data)
        if self._history_manager.execute(command):
            # We don't need to full reload, just refresh view?
            # Ideally we just repaint, but to be safe and consistent:
            # Clear cache for this page
            self._viewer._page_cache.pop(page_num, None)
            self._viewer._render_worker.request_page(
                page_num, self._viewer._zoom)
            self._is_modified = True
            self._update_title()
            self._statusbar.showMessage(f"Added {annot_type} annotation", 2000)

    # ==================== Helpers ====================

    def _update_title(self):
        """Update window title"""
        title = config.APP_NAME
        if self._current_file:
            title = f"{self._current_file.name} - {title}"
        if self._is_modified:
            title = f"*{title}"
        self.setWindowTitle(title)

    def _update_actions_state(self):
        """Update enabled state of actions"""
        has_doc = self._document.is_open

        self._save_action.setEnabled(has_doc)
        self._save_as_action.setEnabled(has_doc)
        self._close_action.setEnabled(has_doc)
        self._print_action.setEnabled(has_doc)
        self._properties_action.setEnabled(has_doc)
        self._find_action.setEnabled(has_doc)

    def _update_recent_files_menu(self):
        """Update recent files menu"""
        self._recent_menu.clear()

        for filepath in self._settings.recent_files[:10]:
            if os.path.exists(filepath):
                action = self._recent_menu.addAction(
                    os.path.basename(filepath),
                    lambda f=filepath: self._open_file(f)
                )
                action.setToolTip(filepath)

        if self._settings.recent_files:
            self._recent_menu.addSeparator()
            self._recent_menu.addAction(
                "Clear Recent Files", self._clear_recent_files)

    def _clear_recent_files(self):
        """Clear recent files list"""
        self._settings.clear_recent_files()
        self._update_recent_files_menu()

    def _format_size(self, size_bytes: int) -> str:
        """Format file size for display"""
        size: float = float(size_bytes)
        for unit in ['B', 'KB', 'MB', 'GB']:
            if size < 1024:
                return f"{size:.1f} {unit}"
            size /= 1024
        return f"{size:.1f} TB"

    def closeEvent(self, event: QCloseEvent):
        """Handle window close"""
        if self._confirm_close():
            self._save_settings()
            self._viewer.cleanup()  # Stop background render thread
            self._document.close()
            event.accept()
        else:
            event.ignore()

    def dragEnterEvent(self, event):
        """Handle drag enter"""
        if event.mimeData().hasUrls():
            for url in event.mimeData().urls():
                if url.toLocalFile().lower().endswith('.pdf'):
                    event.acceptProposedAction()
                    return

    def dropEvent(self, event):
        """Handle file drop"""
        for url in event.mimeData().urls():
            filepath = url.toLocalFile()
            if filepath.lower().endswith('.pdf'):
                self._open_file(filepath)
                break
